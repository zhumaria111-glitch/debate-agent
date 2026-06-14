"""Debate Lens — AI-powered debate analysis.

Paste a B站 video link → transcript + structured mind map + personal knowledge base.
"""

import json
import os
import streamlit as st
from src.video_fetcher import fetch_transcript
from src.transcript_cleaner import clean_transcript
from src.compressor import compress_transcript
from src.dialogue import run_dialogue_turn, generate_welcome_message, get_suggestions
# Lazy-import KnowledgeBase, fixing protobuf compatibility if needed (for Streamlit Cloud Python 3.14).
def _get_kb():
    try:
        from src.knowledge_base import KnowledgeBase
        return KnowledgeBase(persist_dir=os.path.join(ROOT_DIR, "data", "knowledge_base"))
    except Exception:
        import subprocess, sys
        subprocess.check_call([
            sys.executable, "-m", "pip", "install", "--upgrade",
            "protobuf>=5.27.0", "opentelemetry-proto>=1.28.0",
            "opentelemetry-api>=1.28.0", "opentelemetry-sdk>=1.28.0",
            "opentelemetry-exporter-otlp-proto-grpc>=1.28.0", "grpcio>=1.64.0",
        ])
        from src.knowledge_base import KnowledgeBase
        return KnowledgeBase(persist_dir=os.path.join(ROOT_DIR, "data", "knowledge_base"))

# Resolve paths relative to this file so Streamlit works regardless of CWD
ROOT_DIR = os.path.dirname(os.path.abspath(__file__))

# ── Page config ─────────────────────────────────────────────────────────

st.set_page_config(
    page_title="辩析 · Debate Lens",
    page_icon="🎤",
    layout="wide",
    initial_sidebar_state="collapsed",
)

# ── Helpers ─────────────────────────────────────────────────────────────

def _format_duration(seconds: int) -> str:
    """Format seconds to a readable duration string."""
    if seconds <= 0:
        return "—"
    h = seconds // 3600
    m = (seconds % 3600) // 60
    if h > 0:
        return f"{h} 小时 {m} 分钟" if m > 0 else f"{h} 小时"
    return f"{m} 分钟"


def _get_video_duration(url: str) -> int:
    """Quickly fetch video duration from B站 API (fast, <1s). Returns 0 if unknown."""
    import re
    m = re.search(r"BV[a-zA-Z0-9]{10}", url)
    if not m:
        return 0
    try:
        resp = __import__("requests").get(
            f"https://api.bilibili.com/x/web-interface/view?bvid={m.group(0)}",
            headers={"User-Agent": "Mozilla/5.0", "Referer": "https://www.bilibili.com/"},
            timeout=5,
        )
        data = resp.json()
        if data.get("code") == 0:
            dur = int(data["data"].get("duration", 0))
            # Sanity cap: debate videos are < 5 hours (18000 sec)
            return dur if 0 < dur <= 18000 else 0
    except Exception:
        pass
    return 0


# ── Styles ──────────────────────────────────────────────────────────────

st.markdown("""
<style>
    /* Hide default sidebar */
    [data-testid="collapsedControl"] { display: none; }

    /* Video cards */
    .video-card {
        border: 1px solid rgba(128,128,128,0.2);
        border-radius: 12px;
        overflow: hidden;
        cursor: pointer;
        transition: all 0.25s;
        background: rgba(128,128,128,0.05);
        height: 100%;
    }
    .video-card:hover {
        border-color: #ff4b4b;
        box-shadow: 0 4px 16px rgba(255,75,75,0.12);
        transform: translateY(-2px);
    }
    .video-card .cover {
        width: 100%;
        aspect-ratio: 16/10;
        object-fit: cover;
        background: rgba(128,128,128,0.08);
    }
    .video-card .cover-placeholder {
        width: 100%;
        aspect-ratio: 16/10;
        background: linear-gradient(135deg, rgba(45,45,45,0.9) 0%, rgba(26,26,26,0.9) 100%);
        display: flex;
        flex-direction: column;
        align-items: center;
        justify-content: center;
        color: #ccc;
        text-align: center;
        padding: 24px;
    }
    .video-card .cover-placeholder .ph-icon {
        font-size: 32px;
        margin-bottom: 10px;
        opacity: 0.6;
    }
    .video-card .cover-placeholder .ph-title {
        font-size: 14px;
        font-weight: 600;
        line-height: 1.4;
        color: #ddd;
    }
    .video-card .cover-placeholder .ph-tag {
        margin-top: 8px;
        font-size: 10px;
        color: #888;
        border: 1px solid #555;
        padding: 2px 8px;
        border-radius: 3px;
    }
    .video-card .cover-sample {
        width: 100%;
        aspect-ratio: 16/10;
        background: linear-gradient(135deg, rgba(45,45,45,0.9) 0%, rgba(26,26,26,0.9) 100%);
        display: flex;
        flex-direction: column;
        align-items: center;
        justify-content: center;
        color: #ccc;
        text-align: center;
        padding: 24px;
    }
    .video-card .cover-sample .ph-icon {
        font-size: 32px;
        margin-bottom: 10px;
        opacity: 0.6;
    }
    .video-card .cover-sample .ph-title {
        font-size: 14px;
        font-weight: 600;
        line-height: 1.4;
        color: #ddd;
    }
    .video-card .cover-sample .ph-tag {
        margin-top: 8px;
        font-size: 10px;
        color: #888;
        border: 1px solid #555;
        padding: 2px 8px;
        border-radius: 3px;
    }
    .video-card .body {
        padding: 10px 12px 14px 12px;
    }
    .video-card .title {
        font-weight: 600;
        font-size: 14px;
        line-height: 1.4;
        margin-bottom: 6px;
        display: -webkit-box;
        -webkit-line-clamp: 2;
        -webkit-box-orient: vertical;
        overflow: hidden;
    }
    .video-card .meta {
        font-size: 12px;
        opacity: 0.6;
        display: flex;
        gap: 8px;
        align-items: center;
    }
    .video-card .tag {
        background: rgba(255,75,75,0.1);
        color: #ff4b4b;
        padding: 2px 8px;
        border-radius: 4px;
        font-size: 11px;
    }

    /* Feature cards */
    .feature-card {
        border: 1px solid rgba(128,128,128,0.2);
        border-radius: 12px;
        padding: 24px;
        text-align: center;
        background: rgba(128,128,128,0.05);
        height: 100%;
    }
    .feature-card .icon {
        font-size: 36px;
        margin-bottom: 12px;
    }
    .feature-card h3 {
        font-size: 18px;
        margin-bottom: 8px;
    }
    .feature-card p {
        font-size: 13px;
        opacity: 0.7;
        line-height: 1.6;
    }
    /* Hide auto-generated anchor link icons on headings */
    h1 a, h2 a, h3 a, h4 a { display: none; }

    /* Hero text */
    .hero-subtitle { font-size: 16px; opacity: 0.7; margin-bottom: 4px; }
    .hero-hint { font-size: 14px; opacity: 0.5; }

    /* ── Chat layout: fixed bottom input + bubbles ──────────────── */
    /* Fixed chat input — at viewport bottom (full-width for Q&A, overridden in KB) */
    [data-testid="stChatInput"] {
        position: fixed !important;
        bottom: 0 !important;
        left: 0 !important;
        right: 0 !important;
        z-index: 9999 !important;
        padding: 12px 24px 16px 24px !important;
        background: rgba(255,255,255,0.95) !important;
        box-shadow: 0 -1px 0 rgba(0,0,0,0.06), 0 -4px 16px rgba(0,0,0,0.04) !important;
    }
    /* Space below main content so last messages aren't hidden by fixed input */
    section[data-testid="stMain"] {
        padding-bottom: 100px !important;
    }
    /* ── Message bubbles ────────────────────────────────── */
    [data-testid="stChatMessage"] {
        max-width: 75% !important;
        border-radius: 18px !important;
        padding: 10px 18px !important;
        margin-bottom: 8px !important;
    }
    /* AI messages: left-aligned, subtle gray bg */
    div[data-testid="stChatMessage"]:has(div[data-testid="stChatMessageAvatar"]) {
        margin-right: auto !important;
        background: rgba(128,128,128,0.06) !important;
    }
    /* User messages: right-aligned, subtle blue bg */
    div[data-testid="stChatMessage"]:has(div[data-testid="stChatMessageAvatar"][style*="right"]) {
        margin-left: auto !important;
        margin-right: 0 !important;
        background: rgba(59,130,246,0.08) !important;
    }
    /* ── Completion modal ─────────────────────────────── */
    .completion-overlay {
        position: fixed; top: 0; left: 0; right: 0; bottom: 0;
        background: rgba(0,0,0,0.4); z-index: 99999;
        display: flex; align-items: center; justify-content: center;
    }
    .completion-card {
        background: #fff; border-radius: 16px; padding: 36px 44px;
        max-width: 380px; width: 90%; text-align: center;
        box-shadow: 0 8px 32px rgba(0,0,0,0.12);
    }
    .completion-card h3 {
        font-size: 18px; margin: 0 0 24px 0; font-weight: 600; opacity: 0.85;
    }
    .completion-card .stat-row {
        display: flex; justify-content: space-between; align-items: baseline;
        padding: 12px 0; border-bottom: 1px solid rgba(0,0,0,0.06);
    }
    .completion-card .stat-label { font-size: 14px; opacity: 0.55; }
    .completion-card .stat-value { font-size: 22px; font-weight: 600; }
    .completion-card .stat-value.save { color: #22c55e; }
    .completion-card .progress-section { margin-top: 28px; }
    .completion-card .progress-bar {
        width: 100%; height: 6px; background: rgba(0,0,0,0.08); border-radius: 3px;
        overflow: hidden; margin-bottom: 8px;
    }
    .completion-card .progress-fill {
        height: 100%; width: 5%; background: #22c55e; border-radius: 3px;
        animation: progress-grow 60s ease-out forwards;
    }
    @keyframes progress-grow {
        from { width: 5%; }
        to { width: 90%; }
    }
    .completion-card .progress-label { font-size: 12px; opacity: 0.35; }
    .completion-card .close-hint { margin-top: 20px; font-size: 12px; opacity: 0.25; }
</style>
""", unsafe_allow_html=True)

# ── API Key ─────────────────────────────────────────────────────────────

def _load_dotenv(path: str) -> None:
    """Load KEY=VALUE pairs from a .env file into os.environ (no dependencies)."""
    if not os.path.exists(path):
        return
    with open(path) as f:
        for line in f:
            line = line.strip()
            if not line or line.startswith("#") or "=" not in line:
                continue
            key, _, value = line.partition("=")
            key, value = key.strip(), value.strip().strip('"').strip("'")
            if key and key not in os.environ:
                os.environ[key] = value

_load_dotenv(os.path.join(ROOT_DIR, ".env"))

API_KEY = os.environ.get("ANTHROPIC_API_KEY", "")
BASE_URL = os.environ.get("ANTHROPIC_BASE_URL", None)  # CC Switch proxy URL

# ── Session state ───────────────────────────────────────────────────────

defaults = {
    "debate_data": None,
    "processing_done": False,
    "messages": [],
    "transcript": "",
    "video_title": "",
    "video_error": "",
    "current_video_url": "",
    "kb": None,  # KnowledgeBase instance
    "kb_messages": [],  # standalone KB chat messages
    "current_page": "分析",  # "分析" or "知识库"
}
for key, val in defaults.items():
    if key not in st.session_state:
        st.session_state[key] = val

# ── Analysis progress modal (renders at page level so visible during sync analysis) ──
if st.session_state.get("show_analysis_modal"):
    duration_str = _format_duration(st.session_state.get("video_duration", 0))
    st.markdown(f"""
    <div class="completion-overlay">
        <div class="completion-card">
            <h3>正在分析中</h3>
            <div class="stat-row">
                <span class="stat-label">原视频时长</span>
                <span class="stat-value">{duration_str}</span>
            </div>
            <div class="stat-row">
                <span class="stat-label">预计耗时</span>
                <span class="stat-value save">3 分钟以内</span>
            </div>
            <div class="progress-section">
                <div class="progress-bar">
                    <div class="progress-fill"></div>
                </div>
                <div class="progress-label">处理中…</div>
            </div>
        </div>
    </div>
    """, unsafe_allow_html=True)

# Init knowledge base (lazy, non-fatal)
if st.session_state.kb is None:
    try:
        st.session_state.kb = _get_kb()
        st.session_state._kb_available = True
    except Exception as e:
        st.session_state.kb = None
        st.session_state._kb_available = False
        st.session_state._kb_error = str(e)

# ── Built-in debate videos ──────────────────────────────────────────────

BUILT_IN_VIDEOS = [
    {
        "title": "是否应该全面推行四天工作制",
        "url": "",
        "bvid": "",
        "type": "政策辩",
        "transcript_file": os.path.join(ROOT_DIR, "data", "sample_debate.txt"),
    },
    {
        "title": "爱情是/不是人类的必需品",
        "url": "https://www.bilibili.com/video/BV1Zk4y1L7Z7",
        "bvid": "BV1Zk4y1L7Z7",
        "type": "价值辩",
        "transcript_file": os.path.join(ROOT_DIR, "data", "eval", "debate_01.txt"),
    },
    {
        "title": "拐卖妇女儿童应该/不应该买卖同罪",
        "url": "https://www.bilibili.com/video/BV1uY4y1m7hb",
        "bvid": "BV1uY4y1m7hb",
        "type": "政策辩",
        "transcript_file": os.path.join(ROOT_DIR, "data", "eval", "debate_02_cleaned.txt"),
    },
    {
        "title": "网络暴力的根源是恶意还是愚蠢",
        "url": "https://www.bilibili.com/video/BV1N6wLeUEnN",
        "bvid": "BV1N6wLeUEnN",
        "type": "事实辩",
        "transcript_file": os.path.join(ROOT_DIR, "data", "eval", "debate_03_cleaned.txt"),
    },
    {
        "title": "脱口秀还是辩论更能回应世界",
        "url": "https://www.bilibili.com/video/BV1uo4y1f7Ba",
        "bvid": "BV1uo4y1f7Ba",
        "type": "哲理辩",
        "transcript_file": os.path.join(ROOT_DIR, "data", "eval", "debate_04_cleaned.txt"),
    },
    {
        "title": "技术是否中立",
        "url": "https://www.bilibili.com/video/BV1rM4y1D7e6",
        "bvid": "BV1rM4y1D7e6",
        "type": "哲理辩",
        "transcript_file": os.path.join(ROOT_DIR, "data", "eval", "debate_05_cleaned.txt"),
    },
]


@st.cache_data(ttl=86400)
def get_video_cover_b64(bvid: str) -> str:
    """Download B站 cover image server-side, return base64 data URI.

    Downloads once per 24h per video. Serves as embedded data URI so the
    browser never makes a cross-origin request to Bilibili's CDN — no
    Referer leakage and no hotlinking.
    """
    if not bvid:
        return ""

    import base64
    import requests
    try:
        # Step 1: get cover URL from Bilibili's public API
        api_url = f"https://api.bilibili.com/x/web-interface/view?bvid={bvid}"
        api_headers = {
            "User-Agent": "Mozilla/5.0",
            "Referer": "https://www.bilibili.com/",
        }
        resp = requests.get(api_url, headers=api_headers, timeout=8)
        data = resp.json()
        if data["code"] != 0:
            return ""
        pic_url = data["data"]["pic"]

        # Step 2: download the image itself (upgrade http → https)
        pic_url = pic_url.replace("http://", "https://", 1)
        img_resp = requests.get(
            pic_url,
            headers={"User-Agent": api_headers["User-Agent"]},
            timeout=10,
        )
        if img_resp.status_code != 200:
            return ""

        # Step 3: encode as data URI so the browser makes zero external requests
        content_type = img_resp.headers.get("content-type", "image/jpeg")
        b64 = base64.b64encode(img_resp.content).decode()
        return f"data:{content_type};base64,{b64}"
    except Exception:
        return ""


model = "deepseek-chat"

# ── Process transcript ──────────────────────────────────────────────────

def process_transcript(transcript_text: str, source_title: str = ""):
    """Run the compression pipeline and set session state."""
    try:
        debate_data = compress_transcript(transcript_text, API_KEY, model, base_url=BASE_URL)
        st.session_state.debate_data = debate_data.to_dict()
        st.session_state.transcript = transcript_text
        st.session_state.processing_done = True
        st.session_state.video_title = source_title

        welcome = generate_welcome_message(st.session_state.debate_data)
        st.session_state.messages = [
            {"role": "assistant", "content": welcome}
        ]
        st.rerun()
    except Exception as e:
        st.error(f"分析出错：{e}")

# ── Helpers ─────────────────────────────────────────────────────────────

def build_mermaid_mindmap(debate_data: dict) -> str:
    topic = debate_data.get("topic", "辩题")
    aff_side = debate_data.get("affirmative_side", "正方")
    neg_side = debate_data.get("negative_side", "反方")
    qv = debate_data.get("quick_view", {})
    lines = ["```mermaid", "mindmap", f"  root({topic})"]
    lines.append(f"    正方：{aff_side}")
    for i, p in enumerate(qv.get("affirmative_top3", []), 1):
        text = _escape_mermaid(p)
        if text:
            lines.append(f"      论点{i}：{text}")
    lines.append(f"    反方：{neg_side}")
    for i, p in enumerate(qv.get("negative_top3", []), 1):
        text = _escape_mermaid(p)
        if text:
            lines.append(f"      论点{i}：{text}")
    hottest = qv.get("hottest_clash", "")
    if hottest:
        lines.append(f"    🔥交锋")
        lines.append(f"      {_escape_mermaid(hottest)}")
    unresolved = qv.get("key_unresolved", "")
    if unresolved:
        lines.append(f"    ⚠️未回应")
        lines.append(f"      {_escape_mermaid(unresolved)}")
    lines.append("```")
    return "\n".join(lines)

def _escape_mermaid(text: str) -> str:
    return text.replace('"', "'").replace("(", "（").replace(")", "）").replace("[", "【").replace("]", "】")

def _render_mermaid(markdown_code: str, height: int = 900):
    """Render Mermaid markdown code as an interactive diagram."""
    # Strip ```mermaid fences
    src = markdown_code.strip()
    for prefix in ("```mermaid", "```"):
        if src.startswith(prefix):
            src = src[len(prefix):].strip()
    if src.endswith("```"):
        src = src[:-3].strip()
    # Escape for safe embedding
    src_escaped = src.replace("`", "\\`").replace("$", "\\$")
    html = f"""<!DOCTYPE html>
<html>
<head>
<script type="module">
import mermaid from 'https://cdn.jsdelivr.net/npm/mermaid@11/dist/mermaid.esm.min.mjs';
mermaid.initialize({{
    startOnLoad: true,
    theme: 'default',
    mindmap: {{ useMaxWidth: true, padding: 20 }},
    fontSize: 18,
}});
</script>
<style>
body {{ margin: 0; display: flex; justify-content: center; }}
.mermaid svg {{ max-width: 100% !important; height: auto !important; min-height: 700px; }}
</style>
</head>
<body>
<pre class="mermaid">{src_escaped}</pre>
</body>
</html>"""
    st.components.v1.html(html, height=height, scrolling=True)

def build_export_docx(debate_data: dict, messages: list[dict], transcript: str) -> bytes:
    """Generate a Word (.docx) report and return as bytes."""
    from docx import Document
    from docx.shared import Pt, Inches, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    import io

    doc = Document()
    style = doc.styles["Normal"]
    style.font.size = Pt(11)

    # 统一设置中西文字体，避免中英文混排时字体不一致
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement
    rPr = style.element.get_or_add_rPr()
    rFonts = rPr.find(qn("w:rFonts"))
    if rFonts is None:
        rFonts = OxmlElement("w:rFonts")
        rPr.insert(0, rFonts)
    rFonts.set(qn("w:ascii"), "Arial")
    rFonts.set(qn("w:hAnsi"), "Arial")
    rFonts.set(qn("w:eastAsia"), "Arial")
    rFonts.set(qn("w:cs"), "Arial")

    qv = debate_data.get("quick_view") or {}

    # Title
    title = doc.add_heading("辩论分析报告", level=0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # Topic
    doc.add_heading("辩题", level=1)
    doc.add_paragraph(debate_data.get("topic", "未知"))

    # Quick View
    doc.add_heading("三分钟速览", level=1)
    doc.add_paragraph(f"一句话总结：{qv.get('one_sentence_summary') or '（未生成）'}")
    doc.add_paragraph(f"正方立场：{debate_data.get('affirmative_side', '')}")
    doc.add_paragraph(f"反方立场：{debate_data.get('negative_side', '')}")

    doc.add_heading("正方 Top 3 论点", level=2)
    for i, p in enumerate(qv.get("affirmative_top3", []), 1):
        doc.add_paragraph(f"{i}. {p}")

    doc.add_heading("反方 Top 3 论点", level=2)
    for i, p in enumerate(qv.get("negative_top3", []), 1):
        doc.add_paragraph(f"{i}. {p}")

    doc.add_paragraph(f"[最激烈交锋] {qv.get('hottest_clash') or '（未生成）'}")
    doc.add_paragraph(f"[关键未回应问题] {qv.get('key_unresolved') or '（未生成）'}")
    keywords = " | ".join(qv.get("keywords", []))
    if keywords:
        doc.add_paragraph(f"[关键词] {keywords}")

    # Rounds
    rounds = debate_data.get("rounds", [])
    if rounds:
        doc.add_heading("论点攻防链（逐回合）", level=1)
        for r in rounds:
            doc.add_heading(f"{r.get('round_name', '')} — {r.get('side', '')} {r.get('speaker', '')}", level=2)
            doc.add_paragraph(r.get("content_summary", ""))
            for c in r.get("claims", []):
                rebut = f" [反驳 {c.get('rebuts_claim_id')}]" if c.get("rebuts_claim_id") else ""
                status = " [已回应]" if c.get("responded") else " [未回应]"
                p = doc.add_paragraph(f"• {c.get('content', '')}{rebut}{status}")
                if c.get("evidence"):
                    doc.add_paragraph(f"  证据：{c['evidence']}")
                if c.get("response_summary"):
                    doc.add_paragraph(f"  回应：{c['response_summary']}")

    # Unresolved Issues
    unresolved = debate_data.get("unresolved_issues", [])
    if unresolved:
        doc.add_heading("未解决的问题", level=1)
        for u in unresolved:
            doc.add_paragraph(f"[{u.get('importance', '中')}] {u.get('issue', '')}（提出方：{u.get('raised_by', '')}）")

    # Factual Claims
    factual = debate_data.get("factual_claims", [])
    if factual:
        doc.add_heading("事实性主张", level=1)
        for f in factual:
            evidence = f" — 证据：{f.get('evidence_detail')}" if f.get("has_evidence") else " — 无证据"
            doc.add_paragraph(f"• {f.get('claim', '')}（{f.get('speaker', '')}）{evidence}")

    # Q&A
    qa_msgs = [m for m in messages if m["role"] != "assistant" or not m["content"].startswith("这场辩论我已经分析完了")]
    if qa_msgs:
        doc.add_heading("追问记录", level=1)
        for m in qa_msgs:
            role = "AI" if m["role"] == "assistant" else "用户"
            doc.add_paragraph(f"【{role}】{m['content']}")

    # Transcript
    doc.add_heading("辩论逐字稿", level=1)
    doc.add_paragraph(transcript)

    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()


def build_export_report(debate_data: dict, messages: list[dict], transcript: str) -> str:
    """Generate a Markdown report string."""
    qv = debate_data.get("quick_view") or {}
    lines = [f"# 辩论分析报告", "", f"## 辩题", f"{debate_data.get('topic', '未知')}"]

    # ── Quick View ──
    lines.extend(["", "## 速览", "",
        f"**一句话总结**：{qv.get('one_sentence_summary', '（未生成）')}",
        f"**正方立场**：{debate_data.get('affirmative_side', '')}",
        f"**反方立场**：{debate_data.get('negative_side', '')}",
        "", "### 正方 Top 3 论点"])
    for i, p in enumerate(qv.get("affirmative_top3", []), 1):
        lines.append(f"{i}. {p}")
    lines.extend(["", "### 反方 Top 3 论点"])
    for i, p in enumerate(qv.get("negative_top3", []), 1):
        lines.append(f"{i}. {p}")
    lines.extend(["", f"**🔥 最激烈交锋**：{qv.get('hottest_clash') or '（未生成）'}",
        f"**⚠️ 关键未回应问题**：{qv.get('key_unresolved') or '（未生成）'}",
        f"**🏷 关键词**：{' | '.join(qv.get('keywords', []))}"])

    # ── Rounds ──
    rounds = debate_data.get("rounds", [])
    if rounds:
        lines.extend(["", "---", "", "## 论点攻防链（逐回合）", ""])
        for r in rounds:
            lines.append(f"### {r.get('round_name', '')} — {r.get('side', '')} {r.get('speaker', '')}")
            lines.append(f"> {r.get('content_summary', '')}")
            for c in r.get("claims", []):
                rebut = f" [反驳 {c.get('rebuts_claim_id')}]" if c.get("rebuts_claim_id") else ""
                status = " [已回应]" if c.get("responded") else " [未回应]"
                lines.append(f"- **{c.get('content', '')}**{rebut}{status}")
                if c.get("evidence"):
                    lines.append(f"  证据：{c['evidence']}")
                if c.get("response_summary"):
                    lines.append(f"  回应：{c['response_summary']}")
            lines.append("")

    # ── Unresolved Issues ──
    unresolved = debate_data.get("unresolved_issues", [])
    if unresolved:
        lines.extend(["", "## 未解决的问题", ""])
        for u in unresolved:
            lines.append(f"- [{u.get('importance', '中')}] {u.get('issue', '')}（提出方：{u.get('raised_by', '')}）")

    # ── Factual Claims ──
    factual = debate_data.get("factual_claims", [])
    if factual:
        lines.extend(["", "## 事实性主张", ""])
        for f in factual:
            evidence = f" — 证据：{f.get('evidence_detail')}" if f.get("has_evidence") else " — 无证据"
            lines.append(f"- {f.get('claim', '')}（{f.get('speaker', '')}）{evidence}")

    # ── Q&A ──
    qa_msgs = [m for m in messages if m["role"] != "assistant" or not m["content"].startswith("这场辩论我已经分析完了")]
    if qa_msgs:
        lines.extend(["", "---", "", "## 追问记录", ""])
        for m in qa_msgs:
            role = "AI" if m["role"] == "assistant" else "用户"
            lines.append(f"**{role}**：{m['content']}")
            lines.append("")

    # ── Transcript ──
    lines.extend(["", "---", "", "## 辩论逐字稿", "", transcript])
    return "\n".join(lines)

# ── Main UI ─────────────────────────────────────────────────────────────

# --- Header ---
st.markdown("""
<div style="text-align: center; padding: 20px 0 10px 0;">
    <h1 style="font-size: 36px; margin-bottom: 8px;">🎤 辩析 · Debate Lens</h1>
    <p class="hero-subtitle">
        一场一小时辩论赛，三分钟看懂论证骨架
    </p>
    <p class="hero-hint">
        粘贴 B站视频链接即可
    </p>
</div>
""", unsafe_allow_html=True)

# ── Page mode switch ───────────────────────────────────────────────────

current_page = st.radio(
    "模式",
    ["🔍 分析辩论", "🧠 知识库对话"],
    horizontal=True,
    label_visibility="collapsed",
)
# Map pill label back to short key
st.session_state.current_page = "分析" if current_page.startswith("🔍") else "知识库"

# ── 知识库对话 mode (standalone, no per-debate dependency) ──────────────

if st.session_state.current_page == "知识库":
    kb = st.session_state.kb

    if kb is None:
        st.warning("🧠 知识库功能暂不可用。依赖库（ChromaDB/opentelemetry）与当前 Python 版本不兼容。\n\n分析辩论的核心功能不受影响，可以正常使用。")
        st.stop()

    # --- Stats bar ---
    col_kb1, col_kb2, col_kb3 = st.columns([2, 1, 1])
    with col_kb1:
        if kb.count() == 0:
            st.info("知识库为空。先去「分析辩论」分析几场视频，加入知识库后再来。")
        else:
            st.markdown(f"### 🧠 知识库对话 · {kb.count()} 场辩论 · {kb.count_chunks()} 个文本块")
    with col_kb2:
        if st.button("🗑 清空对话", use_container_width=True, key="clear_kb_chat"):
            st.session_state.confirm_clear_kb = True

    if st.session_state.get("confirm_clear_kb"):
        st.warning("确定要清空所有对话记录吗？此操作不可撤销。")
        c1, c2 = st.columns(2)
        with c1:
            if st.button("确认清空", use_container_width=True, key="confirm_clear_kb_yes"):
                st.session_state.kb_messages = []
                st.session_state.confirm_clear_kb = False
                st.rerun()
        with c2:
            if st.button("取消", use_container_width=True, key="confirm_clear_kb_no"):
                st.session_state.confirm_clear_kb = False
                st.rerun()

    if kb.count() == 0:
        st.stop()

    # --- Search + Chat columns ---
    col_search, col_chat = st.columns([2, 3])

    with col_search:
        st.subheader("🔍 检索原文")
        kb_query = st.text_input(
            "搜索知识库",
            placeholder="输入关键词检索所有入库辩论的原文... 如：赌场机制、海德格尔四因说",
            key="kb_search_main",
        )
        if kb_query.strip():
            with st.spinner("检索中..."):
                results = kb.search(kb_query.strip(), n=8)
                transcripts = results["transcripts"]
                if transcripts.get("documents") and transcripts["documents"][0]:
                    st.caption(f"找到 {len(transcripts['documents'][0])} 个相关段落：")
                    for i, doc in enumerate(transcripts["documents"][0]):
                        meta = transcripts["metadatas"][0][i] if transcripts["metadatas"] else {}
                        with st.container(border=True):
                            st.caption(f"📜 {meta.get('topic', '')} · 第{meta.get('chunk_index', 0)+1}/{meta.get('total_chunks', '?')}段")
                            st.markdown(doc[:500] + ("..." if len(doc) > 500 else ""))

                claims = results["claims"]
                if claims.get("documents") and claims["documents"][0]:
                    with st.expander("🎯 相关论点索引"):
                        for i, doc in enumerate(claims["documents"][0]):
                            meta = claims["metadatas"][0][i] if claims["metadatas"] else {}
                            st.caption(f"📌 {meta.get('topic', '')} · {meta.get('side', '')}")
                            st.markdown(f"- {doc}")
                elif not (transcripts.get("documents") and transcripts["documents"][0]):
                    st.caption("未找到相关内容。")

        st.markdown("---")
        st.caption(f"📚 已入库 {kb.count()} 场辩论")
        for d in kb.list_debates():
            col_dn, col_dx = st.columns([5, 1])
            with col_dn:
                st.caption(d["topic"])
            with col_dx:
                if st.button("🗑", key=f"kbmain_del_{d['debate_id']}", help="移除"):
                    kb.remove_debate(d["debate_id"])
                    st.rerun()

    with col_chat:
        st.subheader("💬 AI 研究助手")
        st.caption("基于知识库内容回答，可跨视频对比分析")

        # KB-specific CSS: reposition fixed input to right column, make it visible
        st.markdown("""
        <style>
            [data-testid="stChatInput"] {
                left: 41% !important;
                right: 1.5% !important;
                padding: 12px 0 18px 0 !important;
            }
            [data-testid="stChatInput"] textarea {
                font-size: 15px !important;
                min-height: 48px !important;
                border: 2px solid rgba(128,128,128,0.25) !important;
                border-radius: 12px !important;
                background: rgba(128,128,128,0.03) !important;
            }
            [data-testid="stChatInput"] textarea:focus {
                border-color: rgba(59,130,246,0.5) !important;
                box-shadow: 0 0 0 3px rgba(59,130,246,0.1) !important;
            }
            section[data-testid="stMain"] {
                padding-bottom: 120px !important;
            }
        </style>
        """, unsafe_allow_html=True)

        # Messages
        for msg in st.session_state.kb_messages:
            with st.chat_message(msg["role"]):
                st.markdown(msg["content"])
        # Welcome prompt
        if not st.session_state.kb_messages:
            st.caption("💡 试试这些问题：")
            kb_suggestions = [
                "这知识库里有哪些辩论？各自讨论了什么？",
                "所有辩论中，反方用过哪些类似的论证策略？",
                "帮我对比不同辩论中对「技术」这个概念的定义有什么不同。",
                "哪场辩论里有被回应得最彻底的论点？",
            ]
            cols_ks = st.columns(2)
            for i, sug in enumerate(kb_suggestions):
                with cols_ks[i % 2]:
                    if st.button(sug, key=f"kb_sug_{i}", use_container_width=True):
                        st.session_state.kb_messages.append({"role": "user", "content": sug})
                        st.session_state.needs_kb_response = True
                        st.rerun()

        # Auto-respond to KB preset question
        if st.session_state.get("needs_kb_response"):
            st.session_state.needs_kb_response = False
            kb_input = st.session_state.kb_messages[-1]["content"]
            with st.spinner("检索知识库 + 分析中..."):
                kb_ctx = kb.build_context(kb_input, n=8)
                all_debates = kb.list_debates()
                debate_overview = "\n".join([
                    f"- {d['topic']}：{d['summary'][:100] if d.get('summary') else ''}"
                    for d in all_debates
                ])
                from src.dialogue import run_kb_turn
                response = run_kb_turn(
                    user_message=kb_input,
                    kb_context=kb_ctx,
                    debate_overview=debate_overview,
                    api_key=API_KEY,
                    model=model,
                    base_url=BASE_URL,
                )
            st.session_state.kb_messages.append({"role": "assistant", "content": response})
            st.rerun()

        # Chat input (fixed at bottom)
        if kb_input := st.chat_input("向知识库提问..."):
            st.session_state.kb_messages.append({"role": "user", "content": kb_input})
            with st.spinner("检索知识库 + 分析中..."):
                kb_ctx = kb.build_context(kb_input, n=8)
                all_debates = kb.list_debates()
                debate_overview = "\n".join([
                    f"- {d['topic']}：{d['summary'][:100] if d.get('summary') else ''}"
                    for d in all_debates
                ])
                from src.dialogue import run_kb_turn
                response = run_kb_turn(
                    user_message=kb_input,
                    kb_context=kb_ctx,
                    debate_overview=debate_overview,
                    api_key=API_KEY,
                    model=model,
                    base_url=BASE_URL,
                )
            st.session_state.kb_messages.append({"role": "assistant", "content": response})
            st.rerun()

    st.stop()

# ── Feature cards ──────────────────────────────────────────────────────
_cols_feat = st.columns(3)
with _cols_feat[0]:
    st.markdown("""
    <div class="feature-card">
        <div class="icon">📜</div>
        <h3>一键生成逐字稿</h3>
        <p>粘贴视频链接，自动提取字幕<br>自动清洗标点、识别发言人</p>
    </div>
    """, unsafe_allow_html=True)
with _cols_feat[1]:
    st.markdown("""
    <div class="feature-card">
        <div class="icon">🧠</div>
        <h3>结构化思维导图</h3>
        <p>正反方论点自动提取<br>反驳链、交锋点、未回应问题一目了然</p>
    </div>
    """, unsafe_allow_html=True)
with _cols_feat[2]:
    st.markdown("""
    <div class="feature-card">
        <div class="icon">📥</div>
        <h3>加入个人知识库</h3>
        <p>分析报告 + AI 对话一键导出<br>积累你的辩论学习档案</p>
    </div>
    """, unsafe_allow_html=True)

st.markdown("---")

# ── URL Input ───────────────────────────────────────────────────────────

if not st.session_state.processing_done:
    st.markdown("### 🔗 粘贴视频链接开始分析")

    col_url, col_btn = st.columns([5, 1])
    with col_url:
        video_url = st.text_area(
            "视频链接",
            placeholder="粘贴 B站 / YouTube 视频链接到这里...",
            key="video_url_input",
            label_visibility="collapsed",
            height=68,
        )
    with col_btn:
        analyze_disabled = not (API_KEY and video_url.strip())
        analyze_clicked = st.button(
            "🔍 开始分析",
            type="primary",
            disabled=analyze_disabled,
            use_container_width=True,
            key="analyze_btn",
        )

    if analyze_clicked and video_url.strip():
        st.session_state.pending_video_url = video_url.strip()
        st.session_state.video_duration = _get_video_duration(video_url.strip())
        st.session_state.show_analysis_modal = True
        st.session_state.pending_analysis_url = True
        st.rerun()

    if st.session_state.get("pending_analysis_url"):
        st.session_state.pending_analysis_url = False
        video_url = st.session_state.pop("pending_video_url", "")
        st.session_state.current_video_url = video_url
        st.session_state.processing_done = False
        st.session_state.debate_data = None
        st.session_state.messages = []
        st.session_state.video_error = ""
        st.session_state.video_title = ""

        with st.status("🔍 正在分析辩论...", expanded=True) as status:
            st.write("🎬 获取视频字幕...")
            result = fetch_transcript(video_url)

            if result.error:
                status.update(label="❌ 字幕获取失败", state="error")
                st.session_state.show_analysis_modal = False
                st.error(f"⚠️ 字幕获取失败：{result.error}")
                st.info("💡 该视频可能未开启 AI 字幕。B站视频需开启 AI 字幕才能自动提取文字稿。")
            else:
                st.write("🧹 清洗字幕（加标点、标发言人、去噪声）...")
                cleaned_text = clean_transcript(
                    result.full_text, API_KEY, model, base_url=BASE_URL,
                )
                st.write("🧠 分析论证结构...")
                word_count = len(result.full_text)
                duration = int(getattr(result, "duration", 0) or 0)
                if 0 < duration <= 18000:
                    video_min = duration // 60
                    time_saved = max(1, video_min - 3)
                else:
                    time_saved = max(1, round(word_count / 400) - 3)
                st.session_state.time_saved = time_saved
                st.session_state.word_count = word_count
                st.session_state.video_duration = duration
                st.session_state.show_analysis_modal = False
                status.update(label=f"✅ 分析完成！为你节省了约 {time_saved} 分钟", state="complete")
                process_transcript(cleaned_text, result.title)

    # --- Built-in video gallery ---
    st.markdown("### 📺 热门辩论分析")
    st.caption("以下视频可直接进行分析")

    num_cols = min(3, len(BUILT_IN_VIDEOS))
    for i, vid in enumerate(BUILT_IN_VIDEOS):
        col_idx = i % num_cols
        if col_idx == 0:
            cols_videos = st.columns(num_cols)
        cover_b64 = get_video_cover_b64(vid["bvid"])
        with cols_videos[col_idx]:
            if cover_b64:
                st.markdown(
                    f"""<div class="video-card">
                    <img class="cover" src="{cover_b64}" alt="{vid['title']}"/>
                    <div class="body">
                    <div class="title">{vid['title']}</div>
                    <div class="meta"><span class="tag">{vid['type']}</span>辩论赛</div>
                    </div>
                    </div>""",
                    unsafe_allow_html=True,
                )
            elif not vid["bvid"]:
                # Sample / demo debate — no real video, designed poster cover
                st.markdown(
                    f"""<div class="video-card">
                    <div class="cover-sample">
                        <div class="ph-icon">🧠</div>
                        <div class="ph-title">{vid['title']}</div>
                        <div class="ph-tag">内置范例</div>
                    </div>
                    <div class="body">
                    <div class="title">{vid['title']}</div>
                    <div class="meta"><span class="tag">{vid['type']}</span>辩论赛</div>
                    </div>
                    </div>""",
                    unsafe_allow_html=True,
                )
            else:
                st.markdown(
                    f"""<div class="video-card">
                    <div class="cover-placeholder">📺<br/>{vid['title']}</div>
                    <div class="body">
                    <div class="title">{vid['title']}</div>
                    <div class="meta"><span class="tag">{vid['type']}</span>辩论赛</div>
                    </div>
                    </div>""",
                    unsafe_allow_html=True,
                )

            if st.button(f"分析这场辩论", key=f"builtin_{i}", use_container_width=True):
                st.session_state.pending_builtin_idx = i
                st.session_state.video_duration = _get_video_duration(vid["url"])
                st.session_state.show_analysis_modal = True
                st.session_state.pending_analysis_builtin = True
                st.rerun()

    # Handle pending built-in analysis (outside the loop, after all buttons)
    if st.session_state.get("pending_analysis_builtin"):
        st.session_state.pending_analysis_builtin = False
        idx = st.session_state.pop("pending_builtin_idx", 0)
        vid = BUILT_IN_VIDEOS[idx]
        st.session_state.current_video_url = vid["url"]
        st.session_state.processing_done = False
        st.session_state.debate_data = None
        st.session_state.messages = []
        st.session_state.video_error = ""
        st.session_state.video_title = vid["title"]

        with st.status("📂 正在分析...", expanded=True) as status:
            st.write("📂 加载逐字稿...")
            try:
                with open(vid["transcript_file"], "r") as f:
                    transcript_text = f.read()
                st.write("🧠 分析论证结构...")
                word_count = len(transcript_text)
                video_dur = int(st.session_state.get("video_duration", 0) or 0)
                if 0 < video_dur <= 18000:
                    time_saved = max(1, (video_dur // 60) - 3)
                else:
                    time_saved = max(1, round(word_count / 400) - 3)
                st.session_state.time_saved = time_saved
                st.session_state.word_count = word_count
                st.session_state.show_analysis_modal = False
                status.update(label=f"✅ 完成！节省了约 {time_saved} 分钟", state="complete")
                process_transcript(transcript_text, vid["title"])
            except FileNotFoundError:
                status.update(label="❌ 文件未找到", state="error")
                st.session_state.show_analysis_modal = False
                st.error(f"逐字稿文件未找到：{vid['transcript_file']}")


# ── Results ─────────────────────────────────────────────────────────────

if st.session_state.processing_done and st.session_state.debate_data:
    debate_data = st.session_state.debate_data
    qv = debate_data.get("quick_view", {})
    raw_text = st.session_state.transcript

    st.markdown("---")

    # --- Header ---
    col_title, col_new = st.columns([4, 1])
    with col_title:
        if st.session_state.video_title:
            st.caption(f"📺 {st.session_state.video_title}")
    with col_new:
        if st.button("🔄 分析新的辩论", use_container_width=True, key="new_analysis_top"):
            for key in defaults:
                st.session_state[key] = defaults[key]
            st.rerun()

    # Time-saved badge
    if st.session_state.get("time_saved"):
        video_dur = int(st.session_state.get("video_duration", 0) or 0)
        if 0 < video_dur <= 18000:
            video_min = video_dur // 60
            source_hint = f"原视频 {video_min} 分钟"
        else:
            source_hint = f"原文 {st.session_state.get('word_count', 0):,} 字"
        st.success(
            f"⚡ 分析完成！{source_hint}，"
            f"AI 帮你 3 分钟看懂论证骨架，节省了约 **{st.session_state.time_saved} 分钟**。"
        )

    # --- Function tabs ---
    tab_transcript, tab_mindmap, tab_qa = st.tabs([
        "📜 逐字稿",
        "🧠 思维导图",
        "💬 深度追问",
    ])

    # ── Tab 1: Transcript ───────────────────────────────────────────
    with tab_transcript:
        st.subheader("完整逐字稿")
        if raw_text:
            st.markdown(f"共 {len(raw_text):,} 字")
            st.download_button(
                "📥 下载逐字稿（.txt）",
                data=raw_text,
                file_name=f"辩论逐字稿_{debate_data.get('topic', 'debate')}.txt",
                mime="text/plain",
            )
        else:
            st.info("未获取到逐字稿，可能视频没有字幕。")

    # ── Tab 2: Mind Map ─────────────────────────────────────────────
    with tab_mindmap:
        col_m1, col_m2 = st.columns([3, 1])
        with col_m1:
            st.subheader("论证结构思维导图")
        with col_m2:
            # Word download
            report_docx = build_export_docx(debate_data, st.session_state.messages, raw_text)
            st.download_button(
                "📥 下载详细报告（Word）",
                data=report_docx,
                file_name=f"辩论分析_{debate_data.get('topic', 'report')}.docx",
                mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                use_container_width=True,
            )
            # Add to knowledge base
            kb = st.session_state.kb
            if kb is not None:
                kb_count = kb.count()
                if st.button("🧠 加入知识库", use_container_width=True, help="将本场分析存入知识库，支持跨视频检索"):
                    from src.models import StructuredDebate
                    try:
                        debate_obj = StructuredDebate.from_dict(debate_data)
                        debate_id = kb.add_debate(debate_obj, transcript=raw_text)
                        new_count = kb.count()
                        if new_count > kb_count:
                            st.toast(f"已加入知识库！当前共 {new_count} 场辩论。")
                        else:
                            st.toast("该辩题已在知识库中，无需重复添加。")
                    except Exception as e:
                        st.error(f"加入知识库失败：{e}")

        _render_mermaid(build_mermaid_mindmap(debate_data), height=900)

        st.caption("💡 思维导图展示双方核心论点和交锋结构。切换至「深度追问」tab 可针对任意节点追问。")

        # --- Quick View summary ---
        with st.expander("📋 文字版 3 分钟速览"):
            col_a, col_b = st.columns(2)
            with col_a:
                st.markdown(f"**辩题**：{debate_data.get('topic', '未知')}")
                st.markdown(f"**一句话总结**：{qv.get('one_sentence_summary', '')}")
                st.markdown("**🔥 最激烈交锋**")
                st.info(qv.get("hottest_clash", ""))
                st.markdown("**⚠️ 关键未回应问题**")
                st.warning(qv.get("key_unresolved", ""))
            with col_b:
                st.markdown(f"**正方** · {debate_data.get('affirmative_side', '')}")
                for p in qv.get("affirmative_top3", []):
                    st.markdown(f"- {p}")
                st.markdown(f"**反方** · {debate_data.get('negative_side', '')}")
                for p in qv.get("negative_top3", []):
                    st.markdown(f"- {p}")
                keywords = qv.get("keywords", [])
                kw_html = " ".join(
                    f'<code style="background:#f0f2f6;padding:4px 10px;border-radius:12px;margin:2px;">{kw}</code>'
                    for kw in keywords
                )
                st.markdown(kw_html, unsafe_allow_html=True)

        # --- Argument chain detail ---
        with st.expander("🔍 论点攻防链（逐回合）"):
            rounds = debate_data.get("rounds", [])
            if rounds:
                for r in rounds:
                    st.markdown(f"**{r.get('round_name', '')}** — {r.get('side', '')} {r.get('speaker', '')}")
                    st.caption(r.get("content_summary", ""))
                    for c in r.get("claims", []):
                        rebuts = f" ↳ 反驳 {c.get('rebuts_claim_id')}" if c.get("rebuts_claim_id") else ""
                        status = " ✅已回应" if c.get("responded") else " ⚠️未回应"
                        st.markdown(f"- {c.get('content', '')}{rebuts}{status}")
                    st.markdown("---")

        # --- Unresolved Issues ---
        unresolved = debate_data.get("unresolved_issues", [])
        if unresolved:
            with st.expander("⚠️ 未解决的问题"):
                for u in unresolved:
                    st.markdown(
                        f"**{u.get('issue', '')}** — 提出方：{u.get('raised_by', '')}"
                    )

        # --- Factual Claims ---
        factual = debate_data.get("factual_claims", [])
        if factual:
            with st.expander("📊 事实性主张"):
                for f in factual:
                    evidence = " — 有证据" if f.get("has_evidence") else " — 无证据"
                    st.markdown(f"- {f.get('claim', '')}（{f.get('speaker', '')}）{evidence}")
                    if f.get("evidence_detail"):
                        st.caption(f"  证据详情：{f['evidence_detail']}")

    # ── Tab 3: Deep Q&A ─────────────────────────────────────────────
    with tab_qa:
        # Messages
        for msg in st.session_state.messages:
            with st.chat_message(msg["role"]):
                st.markdown(msg["content"])
        # Suggestions
        user_msg_count = sum(1 for m in st.session_state.messages if m["role"] == "user")
        if user_msg_count == 0:
            st.caption("💡 试试这些问题：")
            suggestions = get_suggestions()
            cols_s = st.columns(len(suggestions))
            for i, sug in enumerate(suggestions):
                with cols_s[i]:
                    if st.button(sug, key=f"sug_{i}", use_container_width=True):
                        st.session_state.messages.append({"role": "user", "content": sug})
                        st.session_state.needs_qa_response = True
                        st.rerun()

        # Auto-respond to preset question
        if st.session_state.get("needs_qa_response"):
            st.session_state.needs_qa_response = False
            user_msg = st.session_state.messages[-1]["content"]
            with st.spinner("🧠 分析中..."):
                history_for_api = [
                    m for m in st.session_state.messages[:-1]
                    if not (
                        m["role"] == "assistant"
                        and m["content"].startswith("这场辩论我已经分析完了")
                    )
                ]
                kb_ctx = ""
                if kb is not None and kb.count() > 0:
                    kb_ctx = kb.build_context(user_msg, n=5)
                response = run_dialogue_turn(
                    user_msg, debate_data, API_KEY, history_for_api,
                    model, base_url=BASE_URL, kb_context=kb_ctx,
                )
            st.session_state.messages.append({"role": "assistant", "content": response})
            st.rerun()

        # Chat input (fixed at bottom)
        if user_input := st.chat_input("追问这场辩论的任何问题..."):
            st.session_state.messages.append({"role": "user", "content": user_input})
            with st.spinner("🧠 分析中..."):
                history_for_api = [
                    m for m in st.session_state.messages[:-1]
                    if not (
                        m["role"] == "assistant"
                        and m["content"].startswith("这场辩论我已经分析完了")
                    )
                ]
                kb_ctx = ""
                if kb is not None and kb.count() > 0:
                    kb_ctx = kb.build_context(user_input, n=5)
                response = run_dialogue_turn(
                    user_input,
                    debate_data,
                    API_KEY,
                    history_for_api,
                    model,
                    base_url=BASE_URL,
                    kb_context=kb_ctx,
                )
            st.session_state.messages.append({"role": "assistant", "content": response})
            st.rerun()

# ── State: No API key ──────────────────────────────────────────────────

if not API_KEY and not st.session_state.processing_done:
    st.warning(
        "⚠️ 未检测到 API Key。请在终端运行：\n\n"
        "```bash\nexport ANTHROPIC_API_KEY=sk-your-key-here\n```\n\n"
        "然后重新启动应用：`streamlit run app.py`"
    )

